"""PDF and Word report generation tests."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from pypdf import PdfReader

from src.models import ReportPayload
from src.reporting.pdf_report import generate_pdf_report
from src.reporting.word_report import generate_word_report


def _payload(frame):
    return ReportPayload(project_title="AI-Powered E-Commerce Analytics", dataset_name="test.csv", dataset_dimensions=f"{len(frame)} rows x {len(frame.columns)} columns", generated_at=datetime.now(timezone.utc), applied_filters={"Region": ["East"]}, question="Which region has the highest sales?", generated_query="SELECT Region, SUM(Sales) FROM dataset GROUP BY Region", query_execution_time_ms=4.2, result_table=frame[["Region", "Sales"]].head(3), narrative="East is the leading region in this verified test result.", key_findings=["The result was computed from filtered data."], limitations="Synthetic test fixture.")


def test_word_report_generation_and_missing_chart(ecommerce_frame):
    content = generate_word_report(_payload(ecommerce_frame))
    assert content.startswith(b"PK")
    document = Document(BytesIO(content))
    assert "AI-Powered E-Commerce Analytics" in "\n".join(p.text for p in document.paragraphs)
    assert len(document.tables) >= 2


def test_pdf_report_generation_and_missing_chart(ecommerce_frame):
    content = generate_pdf_report(_payload(ecommerce_frame))
    assert content.startswith(b"%PDF")
    reader = PdfReader(BytesIO(content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "Which region has the highest sales?" in text
