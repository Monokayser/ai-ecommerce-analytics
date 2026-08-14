"""Polished Word analytics report using the standard business brief preset."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.models import ReportPayload
from src.utils.exceptions import ExportError

NAVY = RGBColor(11, 37, 69)
TEAL = RGBColor(15, 118, 110)
MUTED = RGBColor(100, 116, 139)


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    """Set fixed DXA widths, 120-DXA indent, grid, and matching cell widths."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)


def _field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, end):
        run._r.append(element)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, TEAL),
        ("Heading 2", 13, 12, 6, TEAL),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def generate_word_report(payload: ReportPayload) -> bytes:
    """Generate a styled DOCX report and return its bytes."""
    try:
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(0.492)
        _configure_styles(document)

        header = section.header.paragraphs[0]
        header.text = "AI E-COMMERCE ANALYTICS | VERIFIED QUERY REPORT"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(8)
            run.font.color.rgb = MUTED
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("Page ")
        _field(footer, "PAGE")
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = MUTED

        title = document.add_paragraph()
        title.paragraph_format.space_before = Pt(8)
        title.paragraph_format.space_after = Pt(4)
        run = title.add_run(payload.project_title)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(24)
        run.font.color.rgb = NAVY
        subtitle = document.add_paragraph("AI Assistant Analysis Report")
        subtitle.paragraph_format.space_after = Pt(14)
        subtitle.runs[0].font.size = Pt(13)
        subtitle.runs[0].font.color.rgb = TEAL

        metadata = document.add_table(rows=0, cols=2)
        metadata.style = "Table Grid"
        for label, value in (
            ("Dataset", payload.dataset_name),
            ("Dimensions", payload.dataset_dimensions),
            ("Generated", payload.generated_at.strftime("%Y-%m-%d %H:%M UTC")),
            ("Query time", f"{payload.query_execution_time_ms:.2f} ms"),
            ("Applied filters", ", ".join(f"{k}={v}" for k, v in payload.applied_filters.items()) or "None"),
        ):
            cells = metadata.add_row().cells
            cells[0].text, cells[1].text = label, str(value)
            cells[0].paragraphs[0].runs[0].bold = True
            _set_cell_shading(cells[0], "F2F4F7")
        _set_table_geometry(metadata, [2700, 6660])

        document.add_heading("Question and Validated Query", level=1)
        document.add_paragraph(payload.question)
        query_paragraph = document.add_paragraph()
        query_run = query_paragraph.add_run(payload.generated_query)
        query_run.font.name = "Consolas"
        query_run.font.size = Pt(9)
        query_run.font.color.rgb = NAVY

        document.add_heading("Result", level=1)
        data = payload.result_table.head(50)
        result_table = document.add_table(rows=1, cols=max(len(data.columns), 1))
        result_table.style = "Table Grid"
        if len(data.columns):
            for index, column in enumerate(data.columns):
                result_table.rows[0].cells[index].text = str(column)
            for _, row in data.iterrows():
                cells = result_table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = "" if value is None else str(value)[:120]
        else:
            result_table.rows[0].cells[0].text = "No result columns"
        for cell in result_table.rows[0].cells:
            _set_cell_shading(cell, "DDF3FA")
            for run in cell.paragraphs[0].runs:
                run.bold = True
        column_count = max(len(data.columns), 1)
        base, remainder = divmod(9360, column_count)
        widths = [base + (1 if index < remainder else 0) for index in range(column_count)]
        _set_table_geometry(result_table, widths)
        if len(payload.result_table) > 50:
            document.add_paragraph("Table truncated to the first 50 rows for readability.")

        if payload.chart_image:
            document.add_heading("Visualization", level=1)
            document.add_picture(BytesIO(payload.chart_image), width=Inches(6.25))

        document.add_heading("Analysis", level=1)
        document.add_paragraph(payload.narrative)
        document.add_heading("Key Findings", level=2)
        for finding in payload.key_findings:
            document.add_paragraph(finding, style="List Bullet")
        document.add_heading("Limitations", level=2)
        document.add_paragraph(payload.limitations)
        disclaimer = document.add_paragraph()
        disclaimer.paragraph_format.space_before = Pt(12)
        run = disclaimer.add_run("Important: " + payload.disclaimer)
        run.bold = True
        run.font.color.rgb = RGBColor(122, 90, 0)

        buffer = BytesIO()
        document.save(buffer)
        data_bytes = buffer.getvalue()
        if payload.output_path:
            payload.output_path.parent.mkdir(parents=True, exist_ok=True)
            payload.output_path.write_bytes(data_bytes)
        return data_bytes
    except Exception as exc:
        raise ExportError("Word report generation failed.") from exc
