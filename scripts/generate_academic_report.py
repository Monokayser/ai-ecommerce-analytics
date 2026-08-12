"""Build the final capstone academic report as a styled Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "generated"
ASSETS = OUT / "assets"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "AI_Ecommerce_Analytics_Academic_Report.docx"

NAVY = "071825"
BLUE = "1F5B7A"
TEAL = "16A6A1"
PALE = "EAF4F5"
LIGHT = "F4F6F9"
MID = "D4E2E7"
DARK = RGBColor(19, 42, 56)
WHITE = RGBColor(255, 255, 255)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(86, 111, 124)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_rule(paragraph, color: str = TEAL, size: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.28)
    sec.footer_distance = Inches(0.28)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Title", 25, NAVY, 0, 12),
        ("Heading 1", 17, BLUE, 0, 8),
        ("Heading 2", 12.5, BLUE, 8, 4),
        ("Heading 3", 11, TEAL, 6, 3),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(8.5)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(70, 92, 104)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(6)

    header = sec.header
    p = header.paragraphs[0]
    p.text = "AI-POWERED E-COMMERCE ANALYTICS  |  CAPSTONE REPORT"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_rule(p, TEAL, 7)
    add_page_field(sec.footer.paragraphs[0])


def add_title(doc: Document, number: str, title: str, kicker: str | None = None) -> None:
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(kicker.upper())
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        r.font.letter_spacing = Pt(0.6) if hasattr(r.font, "letter_spacing") else None
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(f"{number}  {title}")
    add_rule(p, TEAL, 8)


def add_h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def add_para(doc: Document, text: str, *, small: bool = False, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if small:
        p.style = doc.styles["Normal"]
        p.paragraph_format.line_spacing = 1.03
        p.paragraph_format.space_after = Pt(3.5)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    if small:
        for r in p.runs:
            r.font.size = Pt(9.2)


def add_bullets(doc: Document, items: list[str], *, compact: bool = True) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.23)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(2 if compact else 4)
        p.paragraph_format.line_spacing = 1.02 if compact else 1.1
        r = p.add_run(item)
        r.font.size = Pt(9.5 if compact else 10.2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], *, font_size: float = 8.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, (cell, text, width) in enumerate(zip(hdr.cells, headers, widths)):
        set_col_width(cell, width)
        set_cell_fill(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(font_size)
                r.font.bold = True
                r.font.color.rgb = WHITE
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell, text, width in zip(row.cells, values, widths):
            set_col_width(cell, width)
            set_cell_fill(cell, LIGHT if row_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(text)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(font_size)
                    r.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def caption(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Figure Caption")


def add_image(doc: Document, filename: str, width: float, caption_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(ASSETS / filename), width=Inches(width))
    caption(doc, caption_text)


def callout(doc: Document, title: str, text: str, color: str = PALE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_col_width(table.cell(0, 0), 6.72)
    cell = table.cell(0, 0)
    set_cell_fill(cell, color)
    set_cell_margins(cell, 110, 140, 110, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(text)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(0)
    for r in p2.runs:
        r.font.size = Pt(9.4)
        r.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    configure_document(doc)

    # Page 1 — Cover
    band = doc.add_table(rows=1, cols=1)
    band.autofit = False
    set_col_width(band.cell(0, 0), 6.72)
    set_cell_fill(band.cell(0, 0), NAVY)
    set_cell_margins(band.cell(0, 0), 360, 300, 360, 300)
    cell = band.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("CAPSTONE PROJECT REPORT")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AI-Powered\nE-Commerce Analytics Platform")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = WHITE
    p = cell.add_paragraph("Secure natural-language analytics, interactive visualization, advanced anomaly detection, and decision-ready reporting")
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(186, 213, 225)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    meta = doc.add_table(rows=6, cols=2)
    meta.autofit = False
    meta.style = "Table Grid"
    details = [
        ("Programme / Module", "MSc Summer 2026 — Data Visualization"),
        ("Lead author", "S. M. Monowar Kayser"),
        ("Student ID", "[INSERT STUDENT ID]"),
        ("Team members", "[INSERT TEAM MEMBER NAMES, IF APPLICABLE]"),
        ("Repository", "github.com/Monokayser/ai-ecommerce-analytics"),
        ("Submission date", "[INSERT FINAL SUBMISSION DATE]"),
    ]
    for i, (key, value) in enumerate(details):
        set_col_width(meta.cell(i, 0), 1.75)
        set_col_width(meta.cell(i, 1), 4.97)
        set_cell_fill(meta.cell(i, 0), PALE)
        set_cell_fill(meta.cell(i, 1), "FFFFFF")
        set_cell_margins(meta.cell(i, 0), 95, 110, 95, 110)
        set_cell_margins(meta.cell(i, 1), 95, 110, 95, 110)
        meta.cell(i, 0).text = key
        meta.cell(i, 1).text = value
        for r in meta.cell(i, 0).paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(BLUE)
        for r in meta.cell(i, 1).paragraphs[0].runs:
            r.font.size = Pt(9); r.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    callout(doc, "Evidence status", "This report distinguishes implemented functionality from verified results. Performance and analytical figures use the deterministic 2,000-row synthetic development dataset (seed 42). The official 5,000+ row dataset and live-provider benchmark remain final verification gates.")

    # Page 2 — Abstract and contents
    page_break(doc)
    add_title(doc, "", "Abstract and Contents", "Research synopsis")
    add_h2(doc, "Abstract")
    add_para(doc, "This capstone presents a production-style analytics platform that converts heterogeneous e-commerce data into verified analytical evidence. The system accepts CSV, JSON, and Parquet files; preserves raw and cleaned views; profiles schema and quality; executes parameterized DuckDB queries; and offers six responsive Streamlit workspaces. Natural-language questions are converted into typed SQL or restricted pandas plans by an adaptive provider-neutral AI layer. Every generated plan is validated with Pydantic and an abstract-syntax-tree allowlist before execution; model text never executes directly. The platform also provides eight visualization families, IQR and Isolation Forest anomaly detection, two-subset comparison, conversation memory, and Word/PDF/PNG/SVG export. Evaluation on the deterministic 2,000-row synthetic dataset produced 72 passing tests, 84% statement coverage, zero Bandit findings, no known dependency vulnerabilities, and warmed DuckDB median query latencies of 14.85–15.80 ms across three representative aggregations. Live LLM benchmark rows remain explicitly not run because no provider key was configured; therefore no model accuracy is claimed. The result is a reproducible, secure, and extensible decision-support environment aligned with the capstone rubric while clearly documenting its evidence limits.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Keywords: ")
    r.bold = True
    p.add_run("e-commerce analytics; natural-language querying; DuckDB; Streamlit; secure AI; anomaly detection; data visualization")
    add_h2(doc, "Contents")
    toc = [
        ("1", "Introduction and problem definition", "3"), ("2", "Requirements and development methodology", "4"),
        ("3", "System architecture", "5"), ("4", "Data and database design", "6"),
        ("5", "Preprocessing and quality management", "7"), ("6", "User-interface design", "8"),
        ("7", "Visualization and interaction design", "9"), ("8", "AI model and natural-language pipeline", "10"),
        ("9", "Security and execution workflow", "11"), ("10", "Anomaly detection", "12"),
        ("11", "Comparative analysis", "13"), ("12", "Implementation and DevOps", "14"),
        ("13", "Testing and evaluation", "15"), ("14", "Performance analysis", "16"),
        ("15", "AI benchmark and reliability", "17"), ("16", "Limitations, ethics, and responsible AI", "18"),
        ("17", "Future work and conclusion", "19"), ("References and disclosure", "", "20"),
    ]
    add_table(doc, ["Section", "Title", "Page"], [[a, b, c] for a, b, c in toc], [0.65, 5.35, 0.72], font_size=7.8)

    # Page 3
    page_break(doc)
    add_title(doc, "1", "Introduction and Problem Definition", "Context")
    add_h2(doc, "1.1 Problem statement")
    add_para(doc, "Retail datasets are often fragmented across spreadsheet and export formats, contain inconsistent headers and data types, and require specialist knowledge before they can answer ordinary business questions. Traditional dashboards provide predetermined views but do not readily translate ad hoc natural-language questions into reproducible calculations. Conversely, unconstrained generative AI can produce plausible but unsafe or unsupported answers. The central problem is therefore not simply query generation: it is the controlled transformation of a user question into validated, executable, and explainable analytical evidence.")
    add_h2(doc, "1.2 Aim and objectives")
    add_para(doc, "The aim was to build a self-contained e-commerce intelligence workspace that remains useful without an API key and becomes more linguistically flexible when a hosted or local model is available. The platform was designed against the official capstone brief [1] and the stricter production requirements supplied for this implementation.")
    add_bullets(doc, [
        "Ingest and validate realistic CSV, JSON, and Parquet datasets while protecting the application from unsafe files and ambiguous schemas.",
        "Support reproducible direct analytics and natural-language querying through a provider-neutral, schema-constrained AI pipeline.",
        "Present results through responsive KPIs and at least eight interactive chart families with consistent accessible styling.",
        "Implement two advanced analytical features: anomaly detection and comparative analysis.",
        "Export decision-ready Word, PDF, PNG, and SVG artifacts with filters, evidence, limitations, and AI-validation disclaimers.",
        "Measure correctness, security, latency, and maintainability without fabricating live-provider results.",
    ], compact=False)
    add_h2(doc, "1.3 Scope and deliverables")
    add_para(doc, "The delivered repository contains a layered Streamlit application, typed contracts, synthetic development data, deployment and security configuration, automated tests, benchmark specifications, Docker support, documentation, this maximum-20-page report, and a 15-minute presentation. External publication and an official-dataset benchmark are deliberately treated as deployment gates rather than inferred achievements.")
    callout(doc, "Research question", "How can a natural-language e-commerce analytics interface improve accessibility while retaining deterministic computation, explicit security controls, reproducibility, and evidence-based communication?")

    # Page 4
    page_break(doc)
    add_title(doc, "2", "Requirements and Development Methodology", "Planning")
    add_h2(doc, "2.1 Requirements traceability")
    add_table(doc, ["Requirement", "Implemented response", "Verification"], [
        ["Data ingestion", "CSV/JSON/Parquet; 200 MB limit; sanitized names; lazy path for >1M rows", "Loader and hardening tests"],
        ["Dashboard", "Six sections, persistent filters, KPIs, responsive theme, empty states", "Streamlit AppTest and browser QA"],
        ["AI assistant", "Offline planner plus Gemini/OpenAI/Ollama adapters; structured plans", "Provider, offline, sandbox tests"],
        ["Secure execution", "Single SELECT/WITH, dataset-only SQL, restricted pandas AST, time/row limits", "Adversarial security tests"],
        ["Visualization", "Eight required chart families plus typed automatic selection", "Constructor smoke tests"],
        ["Advanced analysis", "IQR + Isolation Forest; two-subset KPI comparison", "Calculation and edge-case tests"],
        ["Export", "Word/PDF reports and Plotly PNG/SVG", "Reopen/signature smoke tests"],
        ["Delivery", "README, Docker, CI, security policy, Git history, GitHub release", "Repository audit"],
    ], [1.25, 3.65, 1.82], font_size=7.7)
    caption(doc, "Table 1. Capstone requirements mapped to implementation and verification evidence.")
    add_h2(doc, "2.2 Iterative development methodology")
    add_para(doc, "An incremental engineering method was used. Each vertical slice began with typed contracts and deterministic business logic, then added the Streamlit presentation layer, automated tests, and operational documentation. Security controls were implemented at the same boundary as query generation rather than added after interface completion. This reduced coupling and allowed the offline planner, hosted providers, and direct analytics to share one execution and visualization path.")
    add_table(doc, ["Phase", "Primary activities", "Exit evidence"], [
        ["Discover", "Rubric decomposition; risks; schema and interface contracts", "Traceability matrix"],
        ["Build", "Data layer, query engine, UI, AI adapters, charts, exports", "Runnable vertical slices"],
        ["Harden", "AST allowlists, upload controls, retries, error handling", "Negative security tests"],
        ["Evaluate", "Unit/AppTest/export/benchmark/performance checks", "Measured test artefacts"],
        ["Package", "Docs, Docker, GitHub release, report, presentation", "Reproducible hand-off"],
    ], [0.9, 4.2, 1.62], font_size=8)
    caption(doc, "Table 2. Development workflow and evidence gates.")

    # Page 5
    page_break(doc)
    add_title(doc, "3", "System Architecture", "Layered design")
    add_image(doc, "architecture.png", 6.7, "Figure 1. Implemented layered architecture and controlled natural-language execution path.")
    add_h2(doc, "3.1 Architectural responsibilities")
    add_para(doc, "The Streamlit shell initializes configuration, state, and routing only. UI-independent logic is grouped into data, LLM, visualization, advanced analytics, reporting, and utility modules. Pydantic contracts—such as DatasetBundle, GeneratedQuery, QueryResult, ChartSpec, AnomalyResult, ComparisonResult, and ReportPayload—form stable boundaries between layers. DuckDB provides in-process analytical SQL and direct pandas/Arrow interoperability [3], while Streamlit session state preserves active filters, page selection, dataset context, and five compact conversation records [2].")
    add_h2(doc, "3.2 End-to-end flow")
    add_para(doc, "A validated upload is canonicalized, cleaned, profiled, and registered as the single logical table dataset. Direct dashboard calculations use parameterized query specifications. Natural-language questions pass through a planner, typed-response validation, SQL or pandas security validation, bounded execution, evidence summarization, chart selection, and reporting. Explicit current filters override conversational context. Model output is advisory until the deterministic security and execution stages accept it.")
    callout(doc, "Design principle", "The architecture separates interpretation from computation: AI proposes a typed plan; application code validates, executes, summarizes, and visualizes it.")

    # Page 6
    page_break(doc)
    add_title(doc, "4", "Data and Database Design", "Data engineering")
    add_h2(doc, "4.1 Development dataset")
    add_para(doc, "The repository ships a deterministic synthetic e-commerce CSV generated with seed 42. It contains 2,000 rows and 14 source columns covering transactions, dates, merchandise hierarchy, customer segment, geography, measures, and fulfilment. It is visibly labelled demo data. The application reports official-demo readiness only when at least 5,000 rows are loaded, consistent with the assignment brief [1]. No claim is made that the development dataset represents a real retailer.")
    add_table(doc, ["Field", "Type / role", "Field", "Type / role"], [
        ["Order ID", "text / identifier", "Order Date", "date / time"],
        ["Ship Date", "date / time", "Product Category", "text / dimension"],
        ["Sub-Category", "text / dimension", "Customer Segment", "text / dimension"],
        ["Region", "text / geography", "Country", "text / geography"],
        ["City", "text / geography", "Sales", "decimal / measure"],
        ["Quantity", "integer / measure", "Discount", "decimal / measure"],
        ["Profit", "decimal / measure", "Ship Mode", "text / dimension"],
    ], [1.35, 2.02, 1.35, 2.0], font_size=8)
    caption(doc, "Table 3. Source schema of the synthetic development dataset.")
    add_h2(doc, "4.2 Logical database model")
    add_para(doc, "The application uses a single analytical relation named dataset rather than a transactional server database. Each query receives a fresh DuckDB connection; filters are bound as parameters; output is limited to 1,000 rows by default; and a 10-second interrupt timer bounds execution. The model is intentionally denormalized because the workload is read-heavy, aggregate-oriented, and scoped to one user-supplied analytical file. Raw and cleaned frames remain separately addressable so transformations are auditable and source values are not overwritten.")
    add_h2(doc, "4.3 Canonical schema and aliases")
    add_para(doc, "A configurable alias registry maps common variants to canonical business fields. Exact canonical names take precedence. Ambiguous aliases are reported rather than silently renamed, preventing two source fields from being collapsed. Missing optional columns disable only dependent metrics and charts, maintaining graceful degradation instead of rejecting the entire dataset.")

    # Page 7
    page_break(doc)
    add_title(doc, "5", "Preprocessing and Data-Quality Management", "Reversible transformation")
    add_image(doc, "data_quality.png", 5.9, "Figure 2. Measured quality indicators for the cleaned 2,000-row development dataset.")
    add_h2(doc, "5.1 Cleaning pipeline")
    add_para(doc, "Default cleaning parses Order Date and Ship Date; safely converts Discount, Sales, Quantity, and Profit to numeric types; trims and normalizes categorical text; removes exact duplicate rows; and adds IQR-based outlier markers for numerical measures. pandas provides the tabular transformation layer [4]. Each CleaningAction records the affected field, operation, count, and explanatory detail. Outliers are not deleted or capped because unusual values may be legitimate business events and are required by the advanced analytics task.")
    add_table(doc, ["Measured indicator", "Result", "Interpretation"], [
        ["Raw / cleaned rows", "2,000 / 2,000", "No source records removed"],
        ["Source / cleaned columns", "14 / 18", "Four reversible outlier-marker fields added"],
        ["Missing cells", "0", "No imputation required for the synthetic file"],
        ["Exact duplicates", "0", "Duplicate rule executed; no matches"],
        ["Sales IQR flags", "161 (8.05%)", "Marked only; requires business review"],
        ["Profit IQR flags", "251 (12.55%)", "Marked only; not labelled fraud/error"],
        ["Raw memory", "0.3924 MB", "Development dataset footprint"],
    ], [1.55, 1.35, 3.82], font_size=8)
    caption(doc, "Table 4. Data-quality and preprocessing evidence.")
    add_h2(doc, "5.2 Profiling and LLM-safe context")
    add_para(doc, "The profiler calculates data type, semantic role, nulls, uniqueness, descriptive statistics, bounded samples, and quality notes. Sample cells are truncated and delimited as untrusted quoted data before they enter a prompt. This provides enough context to ground field selection while preventing cell text from becoming instructions or exposing large portions of the dataset.")

    # Page 8
    page_break(doc)
    add_title(doc, "6", "User-Interface and User-Experience Design", "Responsive workspace")
    add_image(doc, "ui_overview.png", 6.65, "Figure 3. Overview workspace showing the accessible navy/teal theme and explicit demo-data status.")
    add_h2(doc, "6.1 Information architecture")
    add_para(doc, "The interface is organized into Overview, Data Exploration, AI Assistant, Advanced Analytics, Data Quality & Performance, and Report Export. The sidebar centralizes upload controls and global dimensions so the active analytical scope is consistent across pages. Dataset, filters, page selection, and the five most recent AI interactions persist in session state. Reset controls clear filters and conversation state without mutating the source dataset.")
    add_h2(doc, "6.2 Accessibility, responsiveness, and interaction")
    add_para(doc, "The visual system uses high-contrast navy surfaces, teal status cues, scalable system fonts, generous touch targets, descriptive labels, tooltips, visible focus states, reduced-motion support, and responsive card grids. KPI cards display concise formatting and preceding-period comparison when a suitable date range exists. Empty states explain missing columns or filters rather than showing broken charts. Subtle CSS transitions provide feedback without blocking interaction, while a lightweight perspective panel supplies a restrained 3D visual motif without introducing a heavy rendering framework.")
    add_h2(doc, "6.3 AI interaction design")
    add_para(doc, "Users can choose Fast, Balanced, or Deep planning modes, see the current provider/fallback status, submit an example question, and review generated query, validation, execution timing, retry state, result summary, chart, and grounded answer. Hidden reasoning is never displayed. The conversation record stores only the question, typed plan, explicit filters, result summary, and final response.")

    # Page 9
    page_break(doc)
    add_title(doc, "7", "Visualization and Interaction Design", "Analytical communication")
    add_h2(doc, "7.1 Visualization portfolio")
    add_table(doc, ["Visualization", "Primary analytical purpose", "Availability / fallback"], [
        ["Time series", "Sales/profit trend and period comparison", "Requires date + numeric measure"],
        ["Choropleth", "Geographic distribution", "Falls back to ranked geographic bars"],
        ["Correlation heatmap", "Pairwise numeric association", "Requires multiple numeric fields"],
        ["Histogram / box", "Distribution and outlier context", "Requires numeric field"],
        ["Sunburst / treemap", "Hierarchical contribution", "Requires two dimensions + measure"],
        ["Grouped / stacked bar", "Category/subgroup comparison", "Requires categorical dimensions"],
        ["Scatter + trend", "Relationship and directional pattern", "Requires two numeric fields"],
        ["Animated time chart", "Change across time and categories", "Requires date, category, measure"],
    ], [1.35, 3.35, 2.02], font_size=8)
    caption(doc, "Table 5. Required visualization families and graceful-degradation rules.")
    add_image(doc, "region_sales.png", 6.15, "Figure 4. Measured regional sales summary rendered with the shared accessible chart theme.")
    add_h2(doc, "7.2 Automatic chart selection and caption validation")
    add_para(doc, "The ChartSelector inspects result cardinality and typed columns. Dates prioritize lines, geographic dimensions prefer maps, two numerics prefer scatter plots, and categorical aggregations prefer bars or hierarchical views. Users may override the recommendation. Plotly supplies interactive hover, zoom, selection, animation, geographic, hierarchical, statistical, and 3D chart primitives [5]. Before display, captions are checked against computed result summaries so model-generated text cannot introduce unsupported values.")

    # Page 10
    page_break(doc)
    add_title(doc, "8", "AI Model and Natural-Language Query Pipeline", "Model engineering")
    add_h2(doc, "8.1 Provider-neutral architecture")
    add_para(doc, "The LLMClient interface exposes generate_query() and format_result(). The preferred hosted path uses Gemini 3.6 Flash configured through the Google GenAI SDK, while OpenAI and Ollama remain optional adapters. Google’s structured-output interface accepts JSON Schema and Pydantic-compatible definitions [7]. A deterministic local planner provides key-free operation for common e-commerce intents and is the automatic fallback when a provider is unavailable. The AI proposes a GeneratedQuery containing interpreted question, language, query, columns, filters, aggregation, chart, and concise rationale.")
    add_table(doc, ["Mode", "Planning behavior", "Narrative behavior", "Use case"], [
        ["Fast", "Low-effort hosted plan or local plan", "Deterministic computed summary", "Live demos and common KPIs"],
        ["Balanced", "Medium-effort structured plan", "Concise evidence-grounded AI text", "Default general analysis"],
        ["Deep", "High-effort structured plan", "High-effort grounded explanation", "Ambiguous multi-step questions"],
    ], [0.85, 2.05, 2.08, 1.74], font_size=8)
    caption(doc, "Table 6. Adaptive AI modes and latency/quality trade-offs.")
    add_h2(doc, "8.2 Prompt construction and response validation")
    add_para(doc, "Prompts include business context, canonical aliases, column types, bounded safe samples, active filters, query rules, and the exact output schema. Data values are placed in explicit untrusted-data delimiters and cannot override system instructions. Provider responses are parsed into Pydantic models; invalid, empty, truncated, or schema-incompatible output fails closed. Authentication, quota, timeout, and unavailable-model failures produce actionable messages and may trigger the deterministic fallback.")
    add_h2(doc, "8.3 Training and evaluation status")
    add_para(doc, "No foundation model was trained or fine-tuned within this project. Consequently, train/validation splits, accuracy/loss curves, and confusion matrices are not applicable to the LLM component. The correct evaluation unit is the generated analytical plan: required fields, operation, filters, execution success, chart suitability, and grounded narrative. Ten benchmark specifications are stored for live-provider evaluation, but all rows remain not_run until a provider key is supplied.")

    # Page 11
    page_break(doc)
    add_title(doc, "9", "Security and Controlled Execution Workflow", "Trust boundary")
    add_h2(doc, "9.1 SQL security policy")
    add_para(doc, "DuckDB SQL is the primary generated language. sqlglot parses the query into an abstract syntax tree before execution. The validator permits exactly one SELECT or WITH statement, restricts relations to the registered dataset table, rejects DDL/DML, multiple statements, table functions, filesystem scans, attachment commands, network-capable extensions, and unsafe expressions, and then enforces the result-row ceiling. Filter values are parameterized rather than interpolated. A fresh connection and 10-second interrupt timer isolate each query.")
    add_h2(doc, "9.2 Restricted pandas fallback")
    add_para(doc, "The pandas route uses a custom AST interpreter with allowlisted operations and attributes. Python eval() and exec() are never called. Imports, dunder access, dangerous built-ins, arbitrary attributes, filesystem and network operations, and mutation of the source frame are denied. This route exists for constrained operations that are inconvenient in SQL, not as a general code-execution facility.")
    add_table(doc, ["Threat", "Control", "Expected outcome"], [
        ["Prompt injection in cells", "Quote, truncate, delimit, schema-constrain", "Cell text remains data"],
        ["DDL/DML or multiple SQL", "Parsed AST statement allowlist", "Rejected before connection"],
        ["Filesystem/network scan", "Table/function denylist", "No external resource access"],
        ["Unsafe Python", "Custom interpreter; no eval/exec/import", "Rejected operation"],
        ["Runaway query", "Fresh connection; 10 s interrupt; row limit", "Bounded execution"],
        ["Hallucinated narrative", "Narrate computed evidence only; caption checks", "Unsupported claims blocked/disclosed"],
        ["Secret exposure", "Environment configuration; no raw-row logs", "Keys and records excluded"],
    ], [1.45, 3.32, 1.95], font_size=7.7)
    caption(doc, "Table 7. Threat model and implemented controls.")
    add_h2(doc, "9.3 Recovery policy")
    add_para(doc, "A sanitized execution error may be returned to the selected provider for exactly one corrected-query attempt. The corrected plan is fully revalidated. Provider transport has one bounded retry for transient faults. Persistent failure produces a safe error or deterministic fallback rather than progressively relaxing policy.")

    # Page 12
    page_break(doc)
    add_title(doc, "10", "Advanced Analytics I — Anomaly Detection", "Task D3")
    add_image(doc, "anomaly_comparison.png", 5.8, "Figure 5. Anomaly rates produced by two complementary methods on Profit.")
    add_h2(doc, "10.1 Algorithms")
    add_para(doc, "The IQR method calculates Q1, Q3, and IQR = Q3 − Q1, then flags observations outside Q1 − 1.5×IQR and Q3 + 1.5×IQR. It is deterministic, transparent, and suited to skewed business measures, but may mark many legitimate tail observations. Isolation Forest constructs random partition trees; observations isolated in fewer splits receive more anomalous scores [8], [9]. The implementation fits 150 trees to the selected numeric target, uses random_state=42, and exposes contamination from 0.001 to 0.5. Users may select target, group, and date controls.")
    add_table(doc, ["Method", "Target / setting", "Flagged rows", "Flagged share"], [
        ["IQR", "Profit; 1.5×IQR", "251", "12.55%"],
        ["Isolation Forest", "Profit; contamination 0.05; 150 trees", "100", "5.00%"],
    ], [1.35, 3.05, 1.2, 1.12], font_size=8.3)
    caption(doc, "Table 8. Measured anomaly results on the synthetic development dataset.")
    add_h2(doc, "10.2 Interpretation and evaluation")
    add_para(doc, "The methods answer different questions. IQR identifies distributional extremes relative to quartiles; Isolation Forest returns a fixed approximate proportion determined by contamination and ranks unusual values through tree isolation. The higher IQR rate does not imply inferior accuracy, because no ground-truth anomaly labels exist. Precision, recall, ROC curves, and a confusion matrix would therefore be invalid. The interface separates facts—thresholds, scores, counts, rows, and grouped plots—from interpretations that require business validation. Statistical unusualness is not evidence of fraud, error, or causation.")
    callout(doc, "ML evaluation boundary", "This is unsupervised, on-demand analysis. The model is fitted to the active subset; there is no reusable trained artefact and no labelled test set. Evaluation reports reproducibility, parameterization, flag counts, and usability—not classification accuracy.")

    # Page 13
    page_break(doc)
    add_title(doc, "11", "Advanced Analytics II — Comparative Analysis", "Task D5")
    add_image(doc, "east_west.png", 5.85, "Figure 6. East–West subset comparison; display scaling is stated in the axis title.")
    add_h2(doc, "11.1 Comparison design")
    add_para(doc, "Users define two subsets using Region, Product Category, Customer Segment, Country, or date period. The service calculates Sales, Profit, Orders, Quantity, average order value (AOV), profit margin, average discount, absolute and percentage differences, and row counts. It handles empty subsets, unequal sample sizes, and zero denominators. The chart, detail table, and narrative are generated from the same ComparisonResult contract.")
    add_table(doc, ["KPI", "East", "West", "West vs East"], [
        ["Sales", "$185,584.52", "$213,841.94", "+15.23%"],
        ["Profit", "$14,924.37", "$22,444.72", "+50.39%"],
        ["Orders", "502", "491", "−2.19%"],
        ["Units", "2,924", "3,069", "+4.96%"],
        ["AOV", "$369.69", "$435.52", "+17.81%"],
        ["Profit margin", "8.04%", "10.50%", "+2.45 pp"],
        ["Average discount", "14.35%", "12.76%", "−1.59 pp"],
        ["Sample rows", "502", "491", "Comparable sizes"],
    ], [1.45, 1.55, 1.55, 2.17], font_size=8)
    caption(doc, "Table 9. Measured East–West comparison on the synthetic development dataset.")
    add_h2(doc, "11.2 Evidence-based interpretation")
    add_para(doc, "West generated 15.23% more sales and 50.39% more profit despite 2.19% fewer orders. Its AOV was 17.81% higher, its margin was 2.45 percentage points higher, and average discount was 1.59 percentage points lower. These are descriptive associations, not causal effects. The similar row counts reduce—without eliminating—sample-size concern; product mix, timing, and customer composition may explain the differences.")

    # Page 14
    page_break(doc)
    add_title(doc, "12", "Implementation Process and DevOps", "Engineering delivery")
    add_h2(doc, "12.1 Technology stack")
    add_table(doc, ["Layer", "Technology", "Role"], [
        ["Presentation", "Streamlit + CSS", "Responsive navigation, state, accessibility, interaction"],
        ["Data", "pandas + PyArrow", "Typed transformations and file interoperability"],
        ["Analytics", "DuckDB", "In-process parameterized analytical SQL"],
        ["AI", "Google GenAI / OpenAI / Ollama", "Provider-neutral structured plan generation"],
        ["Validation", "Pydantic + sqlglot", "Typed responses and SQL AST policy"],
        ["ML", "scikit-learn", "Isolation Forest anomaly detection"],
        ["Visualization", "Plotly + Kaleido", "Interactive charts and PNG/SVG export"],
        ["Reporting", "python-docx + ReportLab", "Word and PDF decision artefacts"],
        ["Quality", "pytest, coverage, Bandit, pip-audit", "Correctness, coverage, security checks"],
        ["Operations", "Docker, GitHub Actions, Streamlit Cloud docs", "Reproducible deployment"],
    ], [1.2, 2.2, 3.32], font_size=7.7)
    caption(doc, "Table 10. Principal technologies and responsibilities.")
    add_h2(doc, "12.2 Repository structure and scalability")
    add_para(doc, "The root app.py performs configuration and routing. Business logic is organized under src/data, src/llm, src/visualization, src/advanced, src/reporting, src/ui, and src/utils. Tests mirror these concerns. Controlled-path and upload loaders enforce a configurable 200 MB limit; files over one million rows use lazy DuckDB scans to avoid loading the full dataset into pandas. Fresh connections, bounded outputs, modular providers, and stateless report payloads support horizontal Streamlit session scaling, subject to external session and storage infrastructure.")
    add_h2(doc, "12.3 Version control and delivery")
    add_para(doc, "The public repository contains source code, assets, configuration, documentation, tests, benchmark definitions, CI configuration, Docker files, a security policy, contribution guidance, citation metadata, and tagged release v1.4.0. Five meaningful commits document release packaging, preview documentation, interface redesign, hardening, and adaptive AI/branding. No artificial history was generated. External cloud deployment was not claimed because no deployment account was available.")
    callout(doc, "Reproducibility", "Target runtime: Python 3.11. The measured workstation used Python 3.12.13, Windows 10 build 19045, AMD64 (4 physical / 8 logical CPUs), and 31.95 GB RAM. Dependencies are pinned and container configuration is provided.")

    # Page 15
    page_break(doc)
    add_title(doc, "13", "Testing and Evaluation", "Verification")
    add_image(doc, "verification_summary.png", 4.95, "Figure 7. Consolidated correctness, coverage, and security verification summary.")
    add_h2(doc, "13.1 Test strategy")
    add_table(doc, ["Test class", "Representative coverage"], [
        ["Data", "Empty/malformed uploads, aliases, conversion, immutability, profiles, duplicates, quality logs"],
        ["Queries", "Aggregations, simultaneous filters, ranges, sorting, top/bottom N, row limits, empty results"],
        ["Security", "DDL/DML, multi-statements, external scans, imports, dangerous attributes/built-ins, injection cells"],
        ["Visualization", "Typed chart selection plus smoke tests for all eight required constructors"],
        ["Advanced", "IQR, Isolation Forest, subset KPIs, zero division, unequal/empty subsets, missing columns"],
        ["Exports", "Word/PDF generation, document reopening, PDF text, PNG/SVG signatures, missing chart"],
        ["Application", "No-key startup, demo banner, navigation, filter reset, conversation reset, empty states"],
    ], [1.22, 5.5], font_size=7.8)
    caption(doc, "Table 11. Test coverage by subsystem.")
    add_h2(doc, "13.2 Measured results")
    add_para(doc, "The final local verification run completed 72 tests successfully with 84% statement coverage. Bandit reported zero unsuppressed findings; pip-audit found no known dependency vulnerabilities; the secret scan was clean; and the headless Streamlit health endpoint returned HTTP 200. Export smoke tests produced valid Word, PDF, PNG, and SVG artefacts. Docker runtime build validation was not performed because Docker was not installed on the evaluation workstation, which remains a documented gate.")
    add_h2(doc, "13.3 Evaluation validity")
    add_para(doc, "Automated tests demonstrate conformity to implemented specifications, not business correctness for all possible retailer datasets. AppTest validates framework behavior but does not replace assistive-technology audits or a full cross-browser device laboratory. Security tests reduce known attack surface but cannot prove the absence of vulnerabilities; dependency and provider policies must be reviewed continuously.")

    # Page 16
    page_break(doc)
    add_title(doc, "14", "Performance Analysis", "Measured latency")
    add_image(doc, "query_performance.png", 6.35, "Figure 8. Warmed median and p95 latency for three representative DuckDB aggregations.")
    add_table(doc, ["Operation", "Median", "p95", "Runs"], [
        ["Sales by Region", "15.7956 ms", "20.6105 ms", "21"],
        ["Profit by Product Category", "14.8455 ms", "20.0162 ms", "21"],
        ["East Sales by Product Category", "15.7186 ms", "20.1737 ms", "21"],
    ], [3.25, 1.2, 1.2, 1.07], font_size=8.3)
    caption(doc, "Table 12. Warmed filtered-aggregation benchmark results.")
    add_h2(doc, "14.1 Loading and profiling")
    add_para(doc, "Measured load time was 304.8952 ms and schema profiling took 71.6971 ms for the 2,000-row, 14-column synthetic CSV. The raw pandas frame occupied 0.3924 MB. Query timing used a warmed connection path followed by 21 measured runs for each operation; median and p95 are reported because they are more informative than a single run.")
    add_h2(doc, "14.2 Interpretation")
    add_para(doc, "All three median latencies are substantially below the capstone target of 500 ms. However, this is not an official-dataset pass: the measured file is smaller than the required 5,000+ row demonstration dataset, hardware differs between environments, and hosted-model latency was not measured. The correct conclusion is that the deterministic local analytical path is fast on the stated workstation and dataset. A final submission should rerun the same benchmark after loading the approved official dataset and record its dimensions and hardware.")

    # Page 17
    page_break(doc)
    add_title(doc, "15", "AI Benchmark, Reliability, and Result Integrity", "Evaluation boundary")
    add_h2(doc, "15.1 Ten-question benchmark design")
    questions = [
        ["1", "Highest-sales region", "SUM", "bar"], ["2", "Top five profitable sub-categories", "SUM + top 5", "bar"],
        ["3", "Monthly sales and profit trend", "SUM + time", "line"], ["4", "Highest segment AOV", "AVG", "bar"],
        ["5", "East vs West sales", "SUM + filter", "bar"], ["6", "Discount–profit association", "correlation", "scatter"],
        ["7", "Loss-making countries", "SUM + <0", "bar"], ["8", "Top ten orders by sales", "top 10", "bar"],
        ["9", "Highest average-profit ship mode", "AVG", "bar"], ["10", "High-discount negative-profit orders", "filter", "scatter"],
    ]
    add_table(doc, ["ID", "Question intent", "Expected operation", "Chart"], questions, [0.48, 3.45, 1.62, 1.17], font_size=7.7)
    caption(doc, "Table 13. Stored live-provider benchmark specification.")
    add_h2(doc, "15.2 Scoring and current status")
    add_para(doc, "The runner evaluates expected columns, operation and filters, formatting compliance, execution success, chart suitability, and evidence-grounded narrative. Accuracy and completeness scores are meaningful only after an actual provider invocation. At report generation time, all ten rows in benchmark_results.csv are not_run because no API key was configured; score cells are blank and execution_success is false by design. The benchmark command replaces these placeholders with measured results only when explicit live-test configuration is present.")
    callout(doc, "No fabricated AI metrics", "The report does not convert deterministic unit-test success into LLM accuracy. Hosted-model accuracy, latency, token usage, and retry rate remain unmeasured. This preserves academic validity and complies with the assignment’s prohibition on fabricated live results.")
    add_h2(doc, "15.3 Reliability mechanisms")
    add_para(doc, "Reliability is improved through schema grounding, provider-native structured outputs, Pydantic validation, AST security checks, parameterized execution, one corrected-query attempt, bounded transport retry, current-filter precedence, five-turn compact memory, evidence-only narratives, numerical caption checks, and deterministic fallback. These mechanisms validate syntax and computation; they do not prove that a user’s question is unambiguous or that the dataset reflects business reality.")

    # Page 18
    page_break(doc)
    add_title(doc, "16", "Limitations, Ethics, Privacy, and Responsible AI", "Critical reflection")
    add_h2(doc, "16.1 Technical and analytical limitations")
    add_bullets(doc, [
        "The bundled 2,000-row synthetic dataset is suitable for development but not a substitute for the official 5,000+ row demonstration dataset or production data diversity.",
        "Hosted-model quality and latency vary with provider, model version, quotas, and prompt complexity; the live benchmark is currently not run.",
        "The deterministic planner covers known e-commerce intents and is intentionally less flexible than a hosted foundation model.",
        "Isolation Forest is univariate in the current implementation; multivariate context and labelled anomaly validation are future work.",
        "Choropleth accuracy depends on recognizable geography; the interface falls back to ranked bars when mapping is unreliable.",
        "Streamlit is appropriate for the capstone scope but enterprise scale requires external authentication, persistent storage, observability, and multi-tenant isolation.",
        "Automated accessibility and security checks do not replace expert manual audits, user studies, or penetration testing.",
    ], compact=False)
    add_h2(doc, "16.2 Privacy and governance")
    add_para(doc, "Uploaded values are processed locally for deterministic analytics. When a hosted provider is selected, only bounded schema context and truncated samples are included. Users must confirm that data is approved for the selected provider and review current provider terms. Sensitive business data should use an appropriately governed paid endpoint or local Ollama deployment. Secrets are read from environment configuration and are not stored in reports, logs, or conversation records.")
    add_h2(doc, "16.3 Responsible use")
    add_para(doc, "The platform presents AI as an interpretation and planning aid, not an autonomous decision-maker. Generated plans are visible; validation and execution status are disclosed; narratives distinguish computed facts from interpretations; anomalies are explicitly non-causal; and reports include limitations and an AI-validation disclaimer. Users remain responsible for checking field definitions, business context, material decisions, and applicable policy.")
    add_h2(doc, "16.4 AI assistance disclosure")
    add_para(doc, "Generative AI assistance was used to help scaffold, implement, test, review, and document the software and academic artefacts. The project author remains responsible for validating the repository, understanding submitted code and claims, rerunning final benchmarks, replacing placeholders, and complying with institutional authorship and citation rules.")

    # Page 19
    page_break(doc)
    add_title(doc, "17", "Future Work and Conclusion", "Research outcome")
    add_h2(doc, "17.1 Future improvements")
    add_table(doc, ["Priority", "Improvement", "Expected contribution"], [
        ["1", "Load the approved 5,000+ row dataset and run the official performance and live-provider benchmark", "Completes final empirical gates"],
        ["2", "Add governed authentication, row-level authorization, encrypted object storage, and audit events", "Supports multi-user deployment"],
        ["3", "Extend anomaly detection to multivariate and seasonal models with labelled expert review", "Improves contextual anomaly quality"],
        ["4", "Add provider telemetry, prompt/version registry, golden datasets, and regression thresholds", "Detects AI quality drift"],
        ["5", "Conduct WCAG 2.2 manual audit and task-based usability study across browsers and mobile devices", "Strengthens accessibility evidence"],
        ["6", "Introduce semantic metric definitions and dataset contracts", "Reduces ambiguity across organizations"],
        ["7", "Deploy to a managed environment with observability, rate limiting, backups, and disaster recovery", "Improves operational resilience"],
    ], [0.62, 3.9, 2.2], font_size=7.8)
    caption(doc, "Table 14. Prioritized research and engineering roadmap.")
    add_h2(doc, "17.2 Conclusion")
    add_para(doc, "The project demonstrates that natural-language analytics can be made more accessible without surrendering control of computation. Its most important contribution is the explicit boundary between interpretation and execution: a local or hosted model proposes a typed plan, while Pydantic, SQL/pandas AST policies, DuckDB, result summarization, and chart validation determine what is accepted and shown. This design supports flexible interaction while remaining useful without an API key.")
    add_para(doc, "Measured local evidence confirms a coherent engineering baseline: real multi-format ingestion, reversible cleaning, six application sections, eight visualization families, two advanced analytical features, exportable reports and charts, 72 passing tests, 84% coverage, clean dependency/security scans, and approximately 15–16 ms warmed query medians on the stated synthetic dataset and workstation. These results should not be generalized beyond the documented conditions. The official dataset, live model benchmark, deployment account, and manual accessibility/security reviews remain transparent final gates.")
    callout(doc, "Final outcome", "A professional, reproducible, provider-neutral analytics platform and evidence pack that satisfy the implemented capstone scope while preserving academic honesty about unmeasured external results.")

    # Page 20
    page_break(doc)
    add_title(doc, "", "References and Submission Disclosure", "Sources")
    refs = [
        "[1] Capstone Project Assignment: AI Analytics, MSc Summer 2026, supplied assessment brief, 2026.",
        "[2] Streamlit, “Session State,” Streamlit Documentation. https://docs.streamlit.io/develop/api-reference/caching-and-state/st.session_state (accessed 12 Aug. 2026).",
        "[3] DuckDB Foundation, “Python API Overview,” DuckDB Documentation. https://duckdb.org/docs/stable/clients/python/overview (accessed 12 Aug. 2026).",
        "[4] pandas Development Team, “pandas 3.0.5 documentation.” https://pandas.pydata.org/docs/ (accessed 12 Aug. 2026).",
        "[5] Plotly, “Plotly Python Open Source Graphing Library.” https://plotly.com/python/ (accessed 12 Aug. 2026).",
        "[6] T. Gorelik, “sqlglot: Python SQL Parser and Transpiler,” GitHub. https://github.com/tobymao/sqlglot (accessed 12 Aug. 2026).",
        "[7] Google AI for Developers, “Structured outputs — Gemini API.” https://ai.google.dev/gemini-api/docs/structured-output (accessed 12 Aug. 2026).",
        "[8] scikit-learn developers, “IsolationForest,” scikit-learn 1.9.0 documentation. https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.IsolationForest.html (accessed 12 Aug. 2026).",
        "[9] F. T. Liu, K. M. Ting, and Z.-H. Zhou, “Isolation Forest,” in Proc. 8th IEEE ICDM, 2008, pp. 413–422, doi:10.1109/ICDM.2008.17.",
        "[10] Pydantic Services Inc., “Pydantic documentation.” https://docs.pydantic.dev/ (accessed 12 Aug. 2026).",
        "[11] OWASP Foundation, “LLM Prompt Injection Prevention Cheat Sheet.” https://cheatsheetseries.owasp.org/cheatsheets/LLM_Prompt_Injection_Prevention_Cheat_Sheet.html (accessed 12 Aug. 2026).",
        "[12] OpenAI, “Responses API.” https://platform.openai.com/docs/api-reference/responses (accessed 12 Aug. 2026).",
        "[13] Ollama, “Structured outputs.” https://docs.ollama.com/capabilities/structured-outputs (accessed 12 Aug. 2026).",
        "[14] ReportLab, “ReportLab PDF Toolkit.” https://www.reportlab.com/docs/reportlab-userguide.pdf (accessed 12 Aug. 2026).",
    ]
    for ref in refs:
        add_para(doc, ref, small=True)
    add_h2(doc, "Submission checklist")
    add_bullets(doc, [
        "Replace student ID, team-member, submission-date, and deployed-URL placeholders.",
        "Load the approved 5,000+ row dataset and rerun pytest, security scans, warmed performance benchmark, and 10-question live-provider benchmark.",
        "Review institutional AI-assistance policy and retain this disclosure or adapt it with supervisor approval.",
        "Confirm the final PDF remains within 20 pages and the 15-minute presentation timing is rehearsed.",
    ])
    add_h2(doc, "Repository and artefact record")
    add_para(doc, "Repository: https://github.com/Monokayser/ai-ecommerce-analytics\nRelease: https://github.com/Monokayser/ai-ecommerce-analytics/releases/tag/v1.4.0\nLive deployment: [INSERT DEPLOYED WEBSITE URL AFTER PUBLICATION]", small=True)

    core = doc.core_properties
    core.title = "AI-Powered E-Commerce Analytics Platform — Academic Report"
    core.subject = "MSc Data Visualization Capstone Project"
    core.author = "S. M. Monowar Kayser"
    core.keywords = "e-commerce analytics, Streamlit, DuckDB, secure AI, anomaly detection"
    core.comments = "Generated from verified repository evidence; editable submission placeholders remain."
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
